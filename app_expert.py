import streamlit as st
import json
import os
import base64
from io import BytesIO
from PIL import Image
import datetime
from typing import Dict, List, Optional
from dotenv import load_dotenv
import requests
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Chargement des variables d'environnement
load_dotenv()

class ChatbotMultimodal:
    def __init__(self):
        # Vérifier si la clé API est configurée
        self.api_key = os.getenv("GROQ_API_KEY")
        if not self.api_key or self.api_key == "votre_clé_api_groq_ici":
            st.error("Veuillez configurer votre clé API Groq dans le fichier .env")
            st.info("Instructions:")
            st.info("1. Allez sur https://console.groq.com/keys")
            st.info("2. Créez une nouvelle clé API")
            st.info("3. Remplacez 'votre_clé_api_groq_ici' dans le fichier .env par votre clé")
            self.api_key = None
        
        self.base_url = "https://api.groq.com/openai/v1"
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
    def process_image_upload(self, image) -> Optional[str]:
        """Traite l'image uploadée et retourne le base64"""
        if image is None:
            return None
        
        try:
            # Convertir l'image en RGB si nécessaire
            if isinstance(image, Image.Image):
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                
                # Sauvegarder temporairement l'image
                buffer = BytesIO()
                image.save(buffer, format='JPEG')
                buffer.seek(0)
                return base64.b64encode(buffer.getvalue()).decode('utf-8')
        except Exception as e:
            st.error(f"Erreur lors du traitement de l'image: {e}")
        
        return None
    
    def get_chat_response(self, message: str, image: Optional[str] = None, history: List = None) -> str:
        """Obtient une réponse du chatbot avec ou sans image"""
        if self.api_key is None:
            return "Veuillez configurer votre clé API Groq dans le fichier .env pour utiliser le chatbot."
        
        try:
            # Préparer les messages pour l'API
            messages = []
            
            # Ajouter l'instruction système pour le français
            messages.append({
                "role": "system",
                "content": "Tu es un assistant IA nommé LaFayette. Tu dois toujours répondre en français, peu importe la langue de la question. Sois poli, professionnel et utile. Si une image est fournie, analyse-la en français."
            })
            
            # Ajouter l'historique de la conversation
            if history:
                for msg in history:
                    if isinstance(msg, dict) and "role" in msg:
                        # Créer une copie du message sans le timestamp
                        clean_msg = {
                            "role": msg["role"],
                            "content": msg["content"]
                        }
                        messages.append(clean_msg)
            
            # Ajouter le nouveau message
            if image:
                # Message avec image
                messages.append({
                    "role": "user",
                    "content": [
                        {"type": "text", "text": message},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{image}"
                            }
                        }
                    ]
                })
                # Utiliser le modèle avec vision (modèle actuel supporté)
                model = "meta-llama/llama-4-scout-17b-16e-instruct"
            else:
                # Message texte seul
                messages.append({
                    "role": "user",
                    "content": message
                })
                # Utiliser le modèle texte
                model = "llama3-70b-8192"
            
            # Préparer la requête
            data = {
                "model": model,
                "messages": messages,
                "max_tokens": 1024,
                "temperature": 0.7,
            }
            
            # Appeler l'API Groq avec requests
            response = requests.post(
                f"{self.base_url}/chat/completions",
                headers=self.headers,
                json=data,
                timeout=30
            )
            
            if response.status_code == 200:
                result = response.json()
                return result["choices"][0]["message"]["content"]
            else:
                return f"Erreur API: {response.status_code} - {response.text}"
            
        except Exception as e:
            return f"Erreur: {str(e)}"

def create_header_footer(canvas, doc):
    """Crée un en-tête et pied de page personnalisés pour le PDF"""
    canvas.saveState()
    
    # En-tête
    canvas.setFont("Helvetica-Bold", 16)
    canvas.setFillColor(HexColor("#1e40af"))
    canvas.drawString(inch, A4[1] - inch, "LaFayette AI Assistant")
    
    # Pied de page
    canvas.setFont("Helvetica", 10)
    canvas.setFillColor(HexColor("#6b7280"))
    canvas.drawString(inch, 0.75 * inch, f"Généré le {datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')}")
    
    canvas.restoreState()

# Configuration de la page Streamlit
st.set_page_config(
    page_title="LaFayette AI - Expert",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded"
)
st.markdown("""
<style>
    /* Import de polices élégantes */
    @import url('https://fonts.googleapis.com/css2?family=Inter:opsz,wght@14..32,300;14..32,400;14..32,500;14..32,600;14..32,700&family=Playfair+Display:wght@400;500;600;700&display=swap');
    
    /* Reset et fond principal */
    .stApp {
        background: linear-gradient(135deg, #f8fafc 0%, #e2e8f0 50%, #f1f5f9 100%);
        font-family: 'Inter', sans-serif;
        color: #1e293b;
    }
    
    /* Logo et titre principal */
    .logo-container {
        text-align: center;
        margin-bottom: 1rem;
        animation: fadeInDown 0.8s ease;
    }
    
    .logo-text {
        font-family: 'Playfair Display', serif;
        font-size: 3.8rem;
        font-weight: 700;
        background: linear-gradient(135deg, #3b82f6, #6366f1, #8b5cf6);
        -webkit-background-clip: text;
        background-clip: text;
        color: transparent;
        letter-spacing: 1px;
        text-shadow: 0 2px 10px rgba(59, 130, 246, 0.2);
    }
    
    .main-header {
        font-family: 'Playfair Display', serif;
        font-size: 2.2rem;
        font-weight: 500;
        text-align: center;
        margin-bottom: 2rem;
        color: #475569;
        letter-spacing: -0.3px;
        border-bottom: 1px solid rgba(59, 130, 246, 0.3);
        display: inline-block;
        padding-bottom: 0.5rem;
        width: auto;
        margin-left: auto;
        margin-right: auto;
    }
    
    /* Animations */
    @keyframes fadeInDown {
        from {
            opacity: 0;
            transform: translateY(-20px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    @keyframes fadeInUp {
        from {
            opacity: 0;
            transform: translateY(20px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    /* Sidebar raffinée */
    [data-testid="stSidebar"] {
        background: rgba(255, 255, 255, 0.9);
        backdrop-filter: blur(16px);
        border-right: 1px solid rgba(59, 130, 246, 0.2);
    }
    
    [data-testid="stSidebar"] .sidebar-section {
        background: rgba(248, 250, 252, 0.8);
        backdrop-filter: blur(12px);
        border-radius: 24px;
        padding: 1.2rem;
        margin-bottom: 1.8rem;
        border: 1px solid rgba(59, 130, 246, 0.3);
        transition: all 0.2s ease;
    }
    
    [data-testid="stSidebar"] .sidebar-section:hover {
        border-color: rgba(59, 130, 246, 0.6);
        box-shadow: 0 8px 20px rgba(59, 130, 246, 0.1);
    }
    
    .section-title {
        font-family: 'Playfair Display', serif;
        font-size: 1.3rem;
        font-weight: 600;
        color: #1e293b;
        margin-bottom: 1rem;
        border-left: 3px solid #3b82f6;
        padding-left: 0.75rem;
    }
    
    /* Cartes de statistiques */
    .stats-card {
        background: linear-gradient(120deg, rgba(59, 130, 246, 0.05), rgba(139, 92, 246, 0.02));
        border-radius: 20px;
        padding: 1rem;
        border: 1px solid rgba(59, 130, 246, 0.2);
        backdrop-filter: blur(4px);
    }
    
    /* Messages chat */
    .chat-message {
        padding: 1.2rem 1.5rem;
        border-radius: 24px;
        margin-bottom: 1.2rem;
        transition: all 0.25s ease;
        animation: fadeInUp 0.4s ease-out;
        backdrop-filter: blur(8px);
    }
    
    .user-message {
        background: rgba(59, 130, 246, 0.06);
        border: 1px solid rgba(59, 130, 246, 0.2);
        border-radius: 24px 24px 8px 24px;
        margin-left: 2rem;
    }
    
    .assistant-message {
        background: rgba(139, 92, 246, 0.04);
        border: 1px solid rgba(139, 92, 246, 0.2);
        border-radius: 24px 24px 24px 8px;
        margin-right: 2rem;
    }
    
    /* Effet de brillance au survol */
    .chat-message:hover {
        transform: translateX(4px);
        border-color: rgba(165, 180, 252, 0.7);
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.2);
    }
    
    /* Zone de saisie */
    .stTextInput > div > div > input {
        background: rgba(255, 255, 255, 0.8);
        border: 1px solid rgba(59, 130, 246, 0.3);
        border-radius: 32px;
        padding: 0.75rem 1.2rem;
        color: #1e293b;
        font-size: 0.95rem;
        transition: all 0.2s;
    }
    
    .stTextInput > div > div > input:focus {
        border-color: #3b82f6;
        box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.1);
        background: rgba(255, 255, 255, 0.9);
    }
    
    /* Boutons */
    .stButton > button {
        background: linear-gradient(95deg, #4f46e5, #7c3aed);
        border: none;
        border-radius: 32px;
        padding: 0.5rem 1.8rem;
        font-weight: 500;
        font-size: 0.9rem;
        transition: all 0.2s;
        box-shadow: 0 2px 8px rgba(79, 70, 229, 0.3);
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 20px rgba(79, 70, 229, 0.4);
        background: linear-gradient(95deg, #6366f1, #8b5cf6);
    }
    
    /* File uploader */
    .stFileUploader > div > div {
        background: rgba(255, 255, 255, 0.7);
        border: 1px dashed rgba(59, 130, 246, 0.4);
        border-radius: 16px;
        padding: 1rem;
        transition: all 0.2s;
        max-width: 100%;
        overflow: hidden;
    }
    
    .stFileUploader > div > div:hover {
        border-color: #3b82f6;
        background: rgba(255, 255, 255, 0.9);
    }
    
    /* Scrollbar fine */
    ::-webkit-scrollbar {
        width: 6px;
    }
    ::-webkit-scrollbar-track {
        background: #f1f5f9;
    }
    ::-webkit-scrollbar-thumb {
        background: #3b82f6;
        border-radius: 4px;
    }
    
    /* Footer */
    .footer-text {
        text-align: center;
        font-size: 0.75rem;
        color: #64748b;
        padding: 1.5rem 0;
        border-top: 1px solid rgba(100, 116, 139, 0.3);
    }
</style>
""", unsafe_allow_html=True)
      

# Initialisation du chatbot
@st.cache_resource
def load_chatbot():
    return ChatbotMultimodal()

chatbot = load_chatbot()

# Logo et titre expert
st.markdown("""
<div class="logo-container">
    <div class="logo-text glow">LaFayette</div>
</div>
<h1 class="main-header">AI Assistant Expert <span class="emoji">Robot</span></h1>
""", unsafe_allow_html=True)

# Sidebar expert visible avec emojis
with st.sidebar:
    st.markdown('<div class="sidebar-section"><h3 class="section-title">Tableau de bord</h3></div>', unsafe_allow_html=True)
    
    # Statistiques
    if 'messages_count' not in st.session_state:
        st.session_state.messages_count = 0
    if 'images_count' not in st.session_state:
        st.session_state.images_count = 0
    
    st.markdown(f"""
    <div class="stats-card">
        <div style="color: #1e293b; font-weight: 600; margin-bottom: 1rem;">
            <span class="emoji">Brain</span> Statistiques
        </div>
        <div style="color: #475569;">
            <span class="emoji">Message</span> Messages: {st.session_state.messages_count}<br>
            <span class="emoji">Image</span> Images: {st.session_state.images_count}<br>
            <span class="emoji">Check</span> Session: Active
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown('<div class="sidebar-section"><h3 class="section-title">Contrôles</h3></div>', unsafe_allow_html=True)
    
    # Bouton pour effacer la conversation
    if st.button("Vider la conversation", type="secondary"):
        st.session_state.messages = []
        st.session_state.messages_count = 0
        st.session_state.images_count = 0
        st.rerun()
    
    st.markdown('<div class="sidebar-section"><h3 class="section-title">Export</h3></div>', unsafe_allow_html=True)
    
    # Export des messages
    if 'messages' in st.session_state and st.session_state.messages:
        # Export JSON
        json_data = json.dumps(st.session_state.messages, indent=2, ensure_ascii=False)
        st.download_button(
            label="JSON",
            data=json_data,
            file_name=f"conversation_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mime="application/json"
        )
        
        # Export TXT
        txt_content = ""
        for msg in st.session_state.messages:
            role = "Utilisateur" if msg["role"] == "user" else "Assistant"
            txt_content += f"{role}:\n{msg['content']}\n\n" + "-"*30 + "\n"
        
        st.download_button(
            label="TXT",
            data=txt_content,
            file_name=f"conversation_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            mime="text/plain"
        )
        
        # Export PDF professionnel
        def create_professional_pdf():
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
            
            # Styles personnalisés
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=TA_CENTER,
                textColor=HexColor("#1e40af"),
                fontName='Helvetica-Bold'
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=12,
                spaceBefore=20,
                textColor=HexColor("#1f2937"),
                fontName='Helvetica-Bold'
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12,
                leading=16,
                textColor=HexColor("#374151"),
                fontName='Helvetica'
            )
            
            story = []
            
            # Page de titre
            story.append(Spacer(1, 2*inch))
            story.append(Paragraph("LaFayette AI Assistant", title_style))
            story.append(Spacer(1, 0.5*inch))
            story.append(Paragraph(f"Conversation du {datetime.datetime.now().strftime('%d/%m/%Y')}", heading_style))
            story.append(Spacer(1, 0.5*inch))
            story.append(Paragraph(f"Messages: {st.session_state.messages_count} | Images: {st.session_state.images_count}", normal_style))
            story.append(PageBreak())
            
            # Contenu de la conversation
            for i, msg in enumerate(st.session_state.messages):
                role = "Utilisateur" if msg["role"] == "user" else "Assistant"
                
                story.append(Paragraph(f"{role}", heading_style))
                
                # Formater le contenu avec retour à la ligne automatique
                content = msg['content'].replace('\n', '<br/>')
                story.append(Paragraph(content, normal_style))
                
                if i < len(st.session_state.messages) - 1:
                    story.append(Spacer(1, 0.3*inch))
            
            # Construire le PDF avec en-tête/pied de page
            doc.build(story, onFirstPage=create_header_footer, onLaterPages=create_header_footer)
            buffer.seek(0)
            return buffer.getvalue()
        
        pdf_data = create_professional_pdf()
        st.download_button(
            label="PDF",
            data=pdf_data,
            file_name=f"conversation_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
            mime="application/pdf"
        )
        
        # Export Word professionnel
        def create_professional_word():
            doc = Document()
            
            # Style du titre
            title = doc.add_heading('🤖 LaFayette AI Assistant', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Informations
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_para.add_run(f'Conversation du {datetime.datetime.now().strftime("%d/%m/%Y")}\n').bold = True
            info_para.add_run(f'Messages: {st.session_state.messages_count} | Images: {st.session_state.images_count}')
            
            # Saut de page
            doc.add_page_break()
            
            # Contenu de la conversation
            for msg in st.session_state.messages:
                role = "Utilisateur" if msg["role"] == "user" else "Assistant"
                
                heading = doc.add_heading(f'{role}', level=2)
                # Appliquer la couleur bleue au titre
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(59, 130, 235)
                
                content_para = doc.add_paragraph(msg['content'])
                # Appliquer la taille de police au contenu
                for run in content_para.runs:
                    run.font.size = Pt(11)
                
                # Ajouter un espacement
                doc.add_paragraph('')
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        
        word_data = create_professional_word()
        st.download_button(
            label="Word",
            data=word_data,
            file_name=f"conversation_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# Zone principale de chat
col1, col2 = st.columns([3, 1])

with col1:
    st.markdown('<h3 class="section-title">Conversation</h3>', unsafe_allow_html=True)
    
    # Initialiser l'historique des messages
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    # Afficher l'historique des messages
    chat_container = st.container()
    with chat_container:
        for message in st.session_state.messages:
            if message["role"] == "user":
                st.markdown(f"""
                <div class="chat-message user-message">
                    <div style="color: #1e40af; font-weight: 600; margin-bottom: 0.5rem; font-family: 'Playfair Display', serif;">
                        <span class="emoji">User</span> Utilisateur
                    </div>
                    <div style="color: #334155; font-weight: 400; line-height: 1.6;">
                        {message['content']}
                    </div>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div class="chat-message assistant-message">
                    <div style="color: #6366f1; font-weight: 600; margin-bottom: 0.5rem; font-family: 'Playfair Display', serif;">
                        <span class="emoji">Robot</span> Assistant
                    </div>
                    <div style="color: #334155; font-weight: 400; line-height: 1.6;">
                        {message['content']}
                    </div>
                </div>
                """, unsafe_allow_html=True)

with col2:
    st.markdown('<h3 class="section-title">Image</h3>', unsafe_allow_html=True)
    
    # Upload d'image
    uploaded_file = st.file_uploader(
        "Téléchargez une image",
        type=['png', 'jpg', 'jpeg'],
        help="Supporte PNG, JPG, JPEG"
    )
    
    # Afficher l'image uploadée
    if uploaded_file is not None:
        image = Image.open(uploaded_file)
        st.image(image, caption="Image uploadée", use_column_width=True)
        
        # Traiter l'image
        image_base64 = chatbot.process_image_upload(image)
        if image_base64:
            st.success("Image traitée avec succès")
        else:
            st.error("Erreur lors du traitement de l'image")
            image_base64 = None
    else:
        image_base64 = None

# Zone de saisie du message
st.markdown('<h3 class="section-title">Saisir un message</h3>', unsafe_allow_html=True)

# Utiliser un formulaire pour gérer l'envoi
with st.form(key="chat_form", clear_on_submit=True):
    col_input1, col_input2 = st.columns([4, 1])
    
    with col_input1:
        user_input = st.text_input(
            "Votre message:",
            placeholder="Tapez votre message ici...",
            key="user_input"
        )
    
    with col_input2:
        send_button = st.form_submit_button("Envoyer", type="primary")

# Gestion de l'envoi du message
if send_button and user_input.strip():
    # Ajouter le message de l'utilisateur
    user_message = {
        "role": "user",
        "content": user_input,
        "timestamp": datetime.datetime.now().isoformat()
    }
    st.session_state.messages.append(user_message)
    st.session_state.messages_count += 1
    
    # Afficher un spinner pendant le traitement
    with st.spinner("L'assistant réfléchit..."):
        # Obtenir la réponse du chatbot
        response = chatbot.get_chat_response(
            user_input, 
            image_base64, 
            st.session_state.messages[:-1]
        )
    
    # Ajouter la réponse de l'assistant
    assistant_message = {
        "role": "assistant",
        "content": response,
        "timestamp": datetime.datetime.now().isoformat()
    }
    st.session_state.messages.append(assistant_message)
    st.session_state.messages_count += 1
    
    # Compter les images
    if image_base64:
        st.session_state.images_count += 1
    
    # Rafraîchir la page pour afficher la nouvelle réponse
    st.rerun()

# Footer expert
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #cbd5e1; padding: 2rem; font-family: "Space Grotesk", sans-serif;'>
    <div style="font-size: 1.3rem; font-weight: 600; margin-bottom: 1rem;">
        <span class="emoji">Robot</span> LaFayette AI Assistant Expert 
        <span class="emoji">Brain</span>
    </div>
    <div style="opacity: 0.8; font-family: 'Inter', sans-serif;">
        <span class="emoji">Message</span> Messages: {} | 
        <span class="emoji">Image</span> Images: {} | 
        <span class="emoji">Check</span> Session Active
    </div>
</div>
""".format(
    st.session_state.messages_count,
    st.session_state.images_count
), unsafe_allow_html=True)
